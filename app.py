import re
import os
import tempfile
from datetime import datetime

import docx
from docx import Document

import numpy as np
import sounddevice as sd
import wave
import soundfile as sf

import whisper
from gtts import gTTS
import pygame


# -----------------------------
# Config
# -----------------------------
INPUT_DOC = "mugilanQp.docx"   # your question paper file
OUTPUT_DOC = "answers.docx"
RECORD_SECONDS = 10            # recording time per question
SAMPLE_RATE = 16000            # Whisper expects 16kHz audio
WHISPER_MODEL = "base"         # "tiny", "base", "small", "medium", "large"

EXCLUDE_SUBSTRS = {
    "answer all questions",
    "name & signature",
    "department of data science",
    "max. marks",
    "time duration",
    "affiliated to",
    "college",
    "batch:", "class:", "subject title:", "semester:",
    "mid term", "reviewer"
}


# -----------------------------
# Audio Helpers
# -----------------------------
def speak_text(text: str):
    """Convert text to speech and play with pygame."""
    if not pygame.mixer.get_init():
        pygame.mixer.init()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as fp:
        temp_mp3 = fp.name
    try:
        gTTS(text=text, lang="en").save(temp_mp3)
        pygame.mixer.music.load(temp_mp3)
        pygame.mixer.music.play()
        clock = pygame.time.Clock()
        while pygame.mixer.music.get_busy():
            clock.tick(10)
    finally:
        try:
            pygame.mixer.music.unload()
        except Exception:
            pass
        if os.path.exists(temp_mp3):
            os.remove(temp_mp3)


def record_wav(path: str, seconds: int = RECORD_SECONDS, sr: int = SAMPLE_RATE):
    print(f"🎤 Recording for {seconds} seconds... Answer now!")
    audio = sd.rec(int(seconds * sr), samplerate=sr, channels=1, dtype="int16")
    sd.wait()
    with wave.open(path, "w") as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)  # int16
        wf.setframerate(sr)
        wf.writeframes(audio.tobytes())


def transcribe_wav(path: str, model) -> str:
    """Transcribe WAV file with Whisper (no ffmpeg)."""
    audio, sr = sf.read(path, dtype="float32")
    if sr != SAMPLE_RATE:
        raise ValueError(f"Recording must be {SAMPLE_RATE} Hz, got {sr}")
    result = model.transcribe(audio, fp16=False)
    return (result.get("text") or "").strip()


# -----------------------------
# Docx Extractor
# -----------------------------
def get_all_text(doc_path):
    """Extract text from both paragraphs and tables."""
    doc = docx.Document(doc_path)
    lines = []

    # paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    lines.append(txt)

    return lines


def clean_line(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip())


def is_excluded(line: str) -> bool:
    l = line.lower()
    return any(key in l for key in EXCLUDE_SUBSTRS)


def extract_questions(path: str):
    """Extract questions from both paragraphs + tables (handles A/B/C)."""
    raw_lines = get_all_text(path)
    lines, prev = [], None
    for l in raw_lines:
        l = clean_line(l)
        if l and l != prev:   # skip duplicates
            lines.append(l)
        prev = l

    section, result, i = None, [], 0

    while i < len(lines):
        line = lines[i]

        # Detect section headers
        if line.lower().startswith("section a"):
            section = "A"; i += 1; continue
        if line.lower().startswith("section b"):
            section = "B"; i += 1; continue
        if line.lower().startswith("section c"):
            section = "C"; i += 1; continue

        if not section or is_excluded(line):
            i += 1
            continue

        # ----- Section A (MCQs) -----
        if section == "A":
            # Case 1: number + text on same line
            m = re.match(r"^(\d{1,2})\s+(.*)$", line)
            if m:
                qnum, stem_text = m.groups()
                i += 1
            # Case 2: number alone on one line
            elif re.fullmatch(r"\d{1,2}", line):
                qnum = line
                i += 1
                stem_text = ""
                if i < len(lines) and not re.match(r"^[ABCD]\b", lines[i], re.I):
                    stem_text = lines[i].strip()
                    i += 1
            else:
                i += 1
                continue

            # Collect options like "A Data reduction"
            options = {}
            while i < len(lines) and re.match(r"^[ABCD]\b", lines[i], re.I):
                letter = lines[i][0].upper()
                value = lines[i][1:].strip()
                options[letter] = value
                i += 1

            # Build final question text
            text = stem_text
            for letter in ["A", "B", "C", "D"]:
                if letter in options:
                    text += f"\n{letter}. {options[letter]}"

            result.append({"section": "A", "label": qnum, "text": text})
            continue

        # ----- Section B/C -----
        m = re.match(r"^(\d{1,2})\s*([AB])?\s*(.*)$", line)
        if section in {"B","C"} and m:
            qnum, ab, rest = m.groups(); ab = ab or ""
            block = [rest] if rest else []; i += 1
            while i < len(lines) and not re.match(r"^\d{1,2}\s*[AB]?", lines[i]) and not lines[i].lower().startswith("section"):
                if not is_excluded(lines[i]) and lines[i] != "(OR)":
                    if lines[i] in {"A","B"} and not block:
                        ab = lines[i]
                    else:
                        block.append(lines[i])
                i += 1
            text = " ".join(block)
            result.append({"section": section, "label": f"{qnum} {ab}".strip(), "text": text})
            continue

        i += 1

    return result




# -----------------------------
# Save Answers to DOCX
# -----------------------------
def save_answers_docx(path: str, qa_items):
    doc = Document()
    doc.add_heading("Answers Document", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    for item in qa_items:
        doc.add_paragraph(f"Q{item['label']}: {item['text']}", style="List Bullet")
        doc.add_paragraph(f"A{item['label']}: {item.get('answer','').strip()}\n")
    doc.save(path)



# -----------------------------
# Main
# -----------------------------
def main():
    if not os.path.exists(INPUT_DOC):
        raise FileNotFoundError(f"Input DOCX not found: {INPUT_DOC}")

    print("📄 Extracting questions...")
    questions = extract_questions(INPUT_DOC)
    if not questions:
        print("❌ No questions found.")
        return

    # 🔑 Sort in natural order (A->B->C) but drop sections later
    section_order = {"A": 1, "B": 2, "C": 3}

    def sort_key(x):
        m = re.match(r"^\s*(\d+)", x["label"])
        if m:
            qnum = int(m.group(1))
        else:
            qnum = 9999
        return (section_order.get(x["section"], 99), qnum, x["label"].strip())

    questions = sorted(questions, key=sort_key)

    print(f"✅ Found {len(questions)} questions.")
    model = whisper.load_model(WHISPER_MODEL)

    qa_items = []
    for idx, q in enumerate(questions, start=1):
        qtext = q["text"]

        # 🚫 No more Section labels
        print(f"\n📢 Question {idx}: {qtext}")
        try:
            speak_text(f"Question {idx}. {qtext}")
        except Exception as e:
            print(f"(Audio playback skipped: {e})")

        temp_wav = os.path.join(tempfile.gettempdir(), f"ans_{idx}.wav")
        record_wav(temp_wav, seconds=RECORD_SECONDS, sr=SAMPLE_RATE)
        try:
            answer = transcribe_wav(temp_wav, model)
        except Exception as e:
            print(f"Transcription error: {e}")
            answer = ""
        finally:
            if os.path.exists(temp_wav):
                os.remove(temp_wav)

        print(f"✅ Transcribed Answer (Q{idx}): {answer}")
        qa_items.append({"label": str(idx), "text": qtext, "answer": answer})

    # save without sections
    save_answers_docx(OUTPUT_DOC, qa_items)
    print(f"\n🎉 All answers saved to: {OUTPUT_DOC}")



if __name__ == "__main__":
    main()
